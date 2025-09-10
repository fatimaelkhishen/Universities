# Install required packages first:
# pip install pdf2docx python-docx transformers torch docx2pdf

from pdf2docx import Converter
from docx import Document
from transformers import MarianMTModel, MarianTokenizer
import torch
from docx2pdf import convert

# ==== CONFIGURATION ====
INPUT_PDF = r"C:\Users\10265480\Downloads\Tunisia 2023 MICS_French.pdf"
TEMP_DOCX = r"C:\Users\10265480\Downloads\temp.docx"
OUTPUT_DOCX = r"C:\Users\10265480\Downloads\Tunisia_2023_MICS_English.docx"
OUTPUT_PDF = r"C:\Users\10265480\Downloads\Tunisia_2023_MICS_English.pdf"

# ==== PDF → DOCX ====
cv = Converter(INPUT_PDF)
cv.convert(TEMP_DOCX, start=0, end=None)  # convert all pages
cv.close()

# ==== Load MarianMT model ====
model_name = "Helsinki-NLP/opus-mt-fr-en"
tokenizer = MarianTokenizer.from_pretrained(model_name)
model = MarianMTModel.from_pretrained(model_name)
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device)

def translate_text(text):
    if not isinstance(text, str) or not text.strip():
        return text
    inputs = tokenizer([text], return_tensors="pt", padding=True, truncation=True).to(device)
    outputs = model.generate(**inputs)
    translated_text = tokenizer.batch_decode(outputs, skip_special_tokens=True)[0]
    return translated_text

# ==== Open DOCX and translate ====
doc = Document(TEMP_DOCX)

print("Translating paragraphs...")
for i, para in enumerate(doc.paragraphs, 1):
    para.text = translate_text(para.text)
    if i % 10 == 0 or i == len(doc.paragraphs):
        print(f"Translated paragraph {i}/{len(doc.paragraphs)}")

print("Translating tables...")
for t_index, table in enumerate(doc.tables, 1):
    for r_index, row in enumerate(table.rows, 1):
        for c_index, cell in enumerate(row.cells, 1):
            cell.text = translate_text(cell.text)
        if r_index % 5 == 0 or r_index == len(table.rows):
            print(f"Table {t_index}, translated row {r_index}/{len(table.rows)}")

# ==== Save translated DOCX ====
doc.save(OUTPUT_DOCX)

# ==== Convert translated DOCX → PDF ====
convert(OUTPUT_DOCX, OUTPUT_PDF)

print(f"\n✅ Translated PDF saved as: {OUTPUT_PDF}")
